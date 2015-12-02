import copy
import os
import sys
from docx import *
import commit
from openpyxl import load_workbook
from openpyxl.cell import *
import re
import logging

from toLogger import ToLogger


def insert_rows(self, row_idx, cnt, above, copy_style, fill_formulae):
    """Inserts new (empty) rows into worksheet at specified row index.

    :param self:
    :param row_idx: Row index specifying where to insert new rows.
    :param cnt: Number of rows to insert.
    :param above: Set True to insert rows above specified row index.
    :param copy_style: Set True if new rows should copy style of immediately above row.
    :param fill_formulae: Set True if new rows should take on formula from immediately above row, filled with references new to rows.

    Usage:

    * insert_rows(2, 10, above=True, copy_style=False)

    """
    CELL_RE = re.compile("(?P<col>\$?[A-Z]+)(?P<row>\$?\d+)")

    row_idx = row_idx - 1 if above else row_idx

    def replace(m):
        row = m.group('row')
        prefix = "$" if row.find("$") != -1 else ""
        row = int(row.replace("$", ""))
        row += cnt if row > row_idx else 0
        return m.group('col') + prefix + str(row)

    # First, we shift all cells down cnt rows...
    old_cells = set()
    old_fas = set()
    new_cells = dict()
    new_fas = dict()
    for c in self._cells.values():

        old_coor = c.coordinate

        # Shift all references to anything below row_idx
        if c.data_type == Cell.TYPE_FORMULA:
            c.value = CELL_RE.sub(
                replace,
                c.value
            )
            # Here, we need to properly update the formula references to reflect new row indices
            if old_coor in self.formula_attributes and 'ref' in self.formula_attributes[old_coor]:
                self.formula_attributes[old_coor]['ref'] = CELL_RE.sub(
                    replace,
                    self.formula_attributes[old_coor]['ref']
                )

        # Do the magic to set up our actual shift
        if c.row > row_idx:
            old_coor = c.coordinate
            old_cells.add((c.row, c.col_idx))
            c.row += cnt
            new_cells[(c.row, c.col_idx)] = c
            if old_coor in self.formula_attributes:
                old_fas.add(old_coor)
                fa = self.formula_attributes[old_coor].copy()
                new_fas[c.coordinate] = fa

    for coor in old_cells:
        del self._cells[coor]
    self._cells.update(new_cells)

    for fa in old_fas:
        del self.formula_attributes[fa]
    self.formula_attributes.update(new_fas)

    # Next, we need to shift all the Row Dimensions below our new rows down by cnt...
    for row in range(len(self.row_dimensions) - 1 + cnt, row_idx + cnt, -1):
        new_rd = copy.copy(self.row_dimensions[row - cnt])
        new_rd.index = row
        self.row_dimensions[row] = new_rd
        del self.row_dimensions[row - cnt]

    # Now, create our new rows, with all the pretty cells
    row_idx += 1
    for row in range(row_idx, row_idx + cnt):
        # Create a Row Dimension for our new row
        new_rd = copy.copy(self.row_dimensions[row - 1])
        new_rd.index = row
        self.row_dimensions[row] = new_rd
        for col in range(1, self.max_column):
            col = get_column_letter(col)
            cell = self.cell('%s%d' % (col, row))
            source = self.cell('%s%d' % (col, row - 1))
            cell.value = source.value
            if copy_style:
                cell.number_format = source.number_format
                cell.font = source.font.copy()
                cell.alignment = source.alignment.copy()
                cell.border = source.border.copy()
                cell.fill = source.fill.copy()
            if fill_formulae and source.data_type == Cell.TYPE_FORMULA:
                s_coor = source.coordinate
                if s_coor in self.formula_attributes and 'ref' not in self.formula_attributes[s_coor]:
                    fa = self.formula_attributes[s_coor].copy()
                    self.formula_attributes[cell.coordinate] = fa

                cell.data_type = Cell.TYPE_FORMULA
                """cell.value = re.sub(
                    "(\$?[A-Z]{1,3}\$?)%d" % (row -1),
                    lambda m: m.group(1) + str(row),
                    source.value
                )"""
    # Check for Merged Cell Ranges that need to be expanded to contain new cells
    for cr_idx, cr in enumerate(self.merged_cell_ranges):
        self.merged_cell_ranges[cr_idx] = CELL_RE.sub(
            replace,
            cr
        )


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def duplicate_rows(table, row, n, c):
    cells = table.add_row().cells
    for i, cell in enumerate(row.cells):
        if '{commit.seq}' in cell.text:
            cells[i].text = cell.text.replace('{commit.seq}', str(n + 1))
            continue
        elif '{commit.module}' in cell.text:
            cells[i].text = cell.text.replace('{commit.module}',
                                              os.path.dirname(os.path.join(c.files[n])).split('/')[0])
            continue
        elif '{commit.file_path}' in cell.text:
            cells[i].text = cell.text.replace('{commit.file_path}', os.path.dirname(os.path.join(c.files[n])))
            continue
        elif '{commit.file_name}' in cell.text:
            cells[i].text = cell.text.replace('{commit.file_name}', os.path.basename(os.path.join(c.files[n])))
            continue
        elif '{commit.mod}' in cell.text:
            cells[i].text = cell.text.replace('{commit.mod}', c.mods[n])
            continue
        else:
            cells[i].text = cell.text


def duplicate_row_times(table, row, n, commit):
    for i in range(n):
        duplicate_rows(table, row, i, commit)
        update_progress(i / n)
    remove_row(table, row)


def duplicate_row_when(doc, n, commit):
    for table in doc.tables:
        for row in table.rows:
            need_rows = False
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text
                    if '{row}' in paragraph_text:
                        need_rows = True
                        paragraph.text = paragraph_text.replace('{row}', '')
                        break
                if need_rows:
                    break
            if need_rows:
                duplicate_row_times(table, row, n, commit)


def change_working_dir():
    # subprocess.call('cd C:/Users/User\PycharmProjects\commit2doc', shell=True)
    os.chdir(os.getcwd())


"""
def getargv():
    version = "1.0.00"
    help = '......(;==)'
    _id = "HEAD"
    if len(sys.argv) >= 2:
        if sys.argv[1].startswith('--'):
            option = sys.argv[1][2:]
            if option == 'version':
                print('Version ' + version)
            elif option == 'help':
                print(help)
            sys.exit()
        else:
            _id = sys.argv[1]
    return _id
"""


def get_dir_list(path):
    lists = []
    for root, subFolders, files in os.walk(path):
        for folder in subFolders:
            lists.append(folder)
        break
    return lists


def update_progress(progress):
    width = 65
    print('\r[{0}{1}] {2}%'.format('#' * int(progress * width), '=' * (width - int(progress * width)),
                                   int(progress * 100)), end='')


def duplicate_row(ws, n):
    for index, row in enumerate(ws.iter_rows()):
        need_copy = False
        for cell in row:
            if cell.value and '{row}' in cell.value:
                need_copy = True
        if need_copy:
            insert_rows(ws, index + 1, n, False, True, True)
            break


def cell_rewrite(ws, search, replace, loop):
    for index, row in enumerate(ws.iter_rows()):
        for cell in row:
            if cell.value and search in cell.value:
                cell.value = cell.value.replace(search, replace)
                if not loop:
                    return


def cell_replace(doc, search, replace):
    searchre = re.compile(search)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text
                    if paragraph_text:
                        if searchre.search(paragraph_text):
                            clear_paragraph(paragraph)
                            paragraph.add_run(paragraph_text.replace(search, replace))


def cell_replace(doc, search, replace):
    searchre = re.compile(search)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text = paragraph.text
                    if paragraph_text:
                        if searchre.search(paragraph_text):
                            clear_paragraph(paragraph)
                            paragraph.add_run(paragraph_text.replace(search, replace))


def clear_paragraph(paragraph):
    p_element = paragraph._p
    p_child_elements = [elm for elm in p_element.iterchildren()]
    for child_element in p_child_elements:
        p_element.remove(child_element)


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)

    return os.path.join(os.path.abspath("."), relative_path)

# ------------------
debug = False

if debug:
    stdout_logger = logging.getLogger('STDOUT')
    sl = ToLogger(stdout_logger, logging.INFO)
    sys.stdout = sl

    stderr_logger = logging.getLogger('STDERR')
    sl = ToLogger(stderr_logger, logging.ERROR)
    sys.stderr = sl

c = commit.Commit('HEAD')
if debug:
    c.dump()

dir_list = get_dir_list(os.getcwd())

input_file = 'input.docx'

for item in dir_list:
    if "SKB" in item:
        input_file = 'skbtpl.docx'
        break
    if "ESUN" in item:
        input_file = 'esuntpl.docx'
        break
    if "YTBK" in item:
        input_file = 'ytbktpl.docx'
        break
    if "IBT" in item:
        input_file = 'ibttpl.xlsx'
        break

if 'xlsx' in input_file:
    wb = load_workbook(os.path.dirname(sys.argv[0]) + '/resource/' + input_file, keep_vba=False)
    print(wb.get_sheet_names())
    ws = wb.get_sheet_by_name(wb.get_sheet_names()[0])
else:
    document = Document(os.path.dirname(sys.argv[0]) + '/resource/' + input_file)

if input_file in 'esuntpl.docx':
    cell_replace(document, '{commit.project_name}', '玉山應收帳款承購管理系統')
    cell_replace(document, '{commit.project_id}', 'ESUNGCL')
if input_file in 'skbtpl.docx':
    cell_replace(document, '{commit.project_name}', '新光應收帳款承購管理系統')
    cell_replace(document, '{commit.project_id}', 'CSPSDB_PRD11222')
if input_file in 'ytbktpl.docx':
    cell_replace(document, '{commit.project_name}', '元大應收帳款承購管理系統')
    cell_replace(document, '{commit.project_id}', 'CSPSDEV1')
if input_file in 'ibttpl.xlsx':
    cell_rewrite(ws, '{commit.project_name}', 'IBT台工銀應收帳款承購管理系統', True)
    cell_rewrite(ws, '{commit.project_id}', 'CSPSDB_IBT156', True)

if '.xlsx' in input_file:
    cell_rewrite(ws, '{commit.author_name}', c.author_name.strip('\n'), True)
    cell_rewrite(ws, '{commit.author_date}', c.author_date.strip('\n'), True)
    cell_rewrite(ws, '{commit.author_email}', ' <' + c.author_email.strip('\n') + '>', True)
    cell_rewrite(ws, '{commit.subject}', '     ' + c.subject.strip('\n'), True)
    cell_rewrite(ws, '{commit.message}', '     ' + c.message.strip('\n'), True)
    cell_rewrite(ws, '{commit.id}', c.id.strip('\n'), True)

    duplicate_row(ws, len(c.files) - 1)

    for num in range(0, len(c.files)):
        cell_rewrite(ws, '{commit.file_name}', os.path.basename(os.path.join(c.files[num])), False)
        cell_rewrite(ws, '{commit.seq}', str(num+1), False)
        cell_rewrite(ws, '{commit.module}', os.path.dirname(os.path.join(c.files[num])).split('/')[0], False)
        cell_rewrite(ws, '{commit.file_path}', os.path.dirname(os.path.join(c.files[num])), False)
        cell_rewrite(ws, '{commit.mod}', c.mods[num], False)
        # cell_rewrite(ws, '{commit._size}', c.file_size[num], False)TODO: get file size

    cell_rewrite(ws, '{row}', '', True)
    wb.save('./commit.xlsx')
else:
    cell_replace(document, '{commit.author_name}', c.author_name.strip('\n'))
    cell_replace(document, '{commit.author_date}', c.author_date.strip('\n'))
    cell_replace(document, '{commit.author_email}', ' <' + c.author_email.strip('\n') + '>')
    cell_replace(document, '{commit.subject}', '     ' + c.subject.strip('\n'))
    cell_replace(document, '{commit.message}', '     ' + c.message.strip('\n'))
    cell_replace(document, '{commit.id}', c.id.strip('\n'))

    duplicate_row_when(document, len(c.files), c)
    document.save('./' + 'commit' + '.docx')

update_progress(1)

print("\nDone....")
